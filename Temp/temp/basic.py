from docx import Document
import requests
from docx.shared import Inches, Pt 
# from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# from pdf import *
# from noteb import *
def get_conclusion(query):
    """Generates a placeholder conclusion related to the problem statement."""
    return f"In conclusion, the problem statement '{query}' has been successfully analyzed and executed."

def replace_text(doc, replacements):
    """Replaces placeholders in a Word document with actual values."""
    
    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                
                para.text = para.text.replace(key, value)
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        
                        cell.text = cell.text.replace(key, value)

def set_cell_color(cell, color):
    cell_xml_element = cell._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), color)
    table_cell_properties.append(shade_obj)

def set_cell_padding(cell, top=100, start=100, bottom=100, end=100):
    """
    Sets cell margins (padding) in twips.
    Default values (100 twips ~ 0.07 inch) can be adjusted.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)

def insert_images(doc, image_replacements):
    """Inserts images into the Word document where required."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, img_path in image_replacements.items():
                    if key in cell.text:
                        cell.text = ""
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run()
                        run.add_picture(img_path, width=Inches(2.5))

def set_table_border(table):
    """Adds a border to the table by modifying its XML."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def fill_template(template_path, output_path, details, num_problems, programs, results):
    doc = Document(template_path)
    
    header_table = doc.add_table(rows=0, cols=2)
    name_row = header_table.add_row().cells
    name_row[0].text = "NAME:"
    name_row[0].paragraphs[0].runs[0].bold = True
    name_row[1].text = "{name}"
    for cell in name_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
    
    uid_row = header_table.add_row().cells
    uid_row[0].text = "UID: "
    uid_row[0].paragraphs[0].runs[0].bold = True
    uid_row[1].text = "{roll_no}"
    for cell in uid_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
    
    exp_row = header_table.add_row().cells
    exp_row[0].text = "EXP NO:"
    exp_row[0].paragraphs[0].runs[0].bold = True
    exp_row[1].text = "{Exp_no}"
    for cell in exp_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
        
    aim_row = header_table.add_row().cells
    aim_row[0].text = "AIM:"
    aim_row[0].paragraphs[0].runs[0].bold = True
    aim_row[1].text = "{aim}"
    for cell in aim_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
    
    set_table_border(header_table)
        
    if num_problems:
        for i in range(1, num_problems + 1):
            # Create the table with no initial rows (we'll add them manually).
            # We use 2 columns.
            table = doc.add_table(rows=0, cols=2)
            
            # Add header row (merged) for Program title.
            header_row = table.add_row().cells
            merged_header = header_row[0].merge(header_row[1])
            merged_header.text = f"\t\t\t\t\t\tPROGRAM {i}"
            # Apply padding and background color.
            set_cell_padding(merged_header, top=150, start=210, bottom=150, end=150)
            set_cell_color(merged_header, 'a6a6a6')
            
            # Add row for PROBLEM STATEMENT.
            prob_row = table.add_row().cells
            prob_row[0].text = "PROBLEM STATEMENT:"
            prob_row[0].paragraphs[0].runs[0].bold = True
            prob_row[1].text = details[f"{{{i}_problem_statement}}"]
            for cell in prob_row:
                set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
            
            #DAA only 
            '''theo_row = table.add_row().cells
            theo_row[0].text = "THEORY:"
            theo_row[0].paragraphs[0].runs[0].bold = True
            theo_row[1].text = details[f"{{program_{i}}}"]
            for cell in theo_row:
                set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
            #DAA Only  
            algo_row = table.add_row().cells
            algo_row[0].text = "Pseudo-Code/Algo:"
            algo_row[0].paragraphs[0].runs[0].bold = True
            algo_row[1].text = details[f"{{program_{i}}}"]
            for cell in algo_row:
                set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
                '''
            
            # Add row for PROGRAM.
            prog_row = table.add_row().cells
            prog_row[0].text = "PROGRAM:"
            prog_row[0].paragraphs[0].runs[0].bold = True
            prog_row[1].text = details[f"{{program_{i}}}"]
            for cell in prog_row:
                set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
            
            # Add row for RESULT (merged across both cells).
            result_row = table.add_row().cells
            merged_result = result_row[0].merge(result_row[1])
            merged_result.text = "RESULT: {{result_%d}}" % i
            if merged_result.paragraphs and merged_result.paragraphs[0].runs:
                merged_result.paragraphs[0].runs[0].bold = True
            set_cell_padding(merged_result, top=150, start=210, bottom=150, end=150)
            
            # DAA Only
            '''an_row = table.add_row().cells
            an_row[0].text = "ANALYSIS:"
            an_row[0].paragraphs[0].runs[0].bold = True
            an_row[1].text = details[f"{{program_{i}}}"]
            for cell in an_row:
                set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
                '''
                
            # Add borders to the table.
            set_table_border(table)
    
    conclusion_table = doc.add_table(rows=0, cols=2)
    concl_row = conclusion_table.add_row().cells
    merged_conclusion = concl_row[0].merge(concl_row[1])
    merged_conclusion.text = "CONCLUSION: " + details["{{conclusion}}"]
    set_cell_padding(merged_conclusion, top=150, start=210, bottom=150, end=150)
    set_table_border(conclusion_table)
    replace_text(doc, details)
    image_placeholders = {f"{{result_{k}}}": v for k, v in results.items()}
    insert_images(doc, image_placeholders)
    
    doc.save(output_path)
    print(f"Document saved as {output_path}")

if __name__ == "__main__":
    template_file = "template.docx"  # Path to your Word template
    output_file = "output.docx"      # Output file
    
    name = input("Enter your name: ")
    roll_no = input("Enter your roll number: ")
    exp_no = input("Enter the experiment number: ")
    aim = input("Enter the aim of the experiment: ")
    num_problems = int(input("Enter the number of problems: "))
    problem_statements = {}
    for i in range(1, num_problems + 1):
        prob = input(f"Enter the problem statement for problem {i}: ")
        problem_statements[i] = prob
    programs = []
    results = {}
    for i in range(1, num_problems + 1):
        print(f"\nFor Program {i}:")
        prog = input("Enter the program code (or a brief description): ")
        programs.append(prog)
        res = input("Enter the path to the result image (or leave blank if none): ")
        if res:
            results[i] = res
            
    
    #------------read the problem statments from the pdf ----------------------------
    # pdf_file_to_read = input("Enter the path to the PDF file: ")
    # extracted_text = read_pdf_text(pdf_file_to_read)
    # if extracted_text is not None:
    #     extracted_text += "\n\n"
    #     pattern = re.compile(r"^\s*(\d{1,2}[\.\)])\s*(.*?)(?=\n\s*\d{1,2}[\.\)]|\n\s*\n\s*\n|\Z)", re.MULTILINE | re.DOTALL)

    #     problem_statements = {}
    #     i = 1
    #     for match in pattern.finditer(extracted_text):
    #         problem_text = match.group(2).strip()
    #         problem_text = re.sub(r'^Problem Statement:\s*', '', problem_text)
    #         problem_statements[i] = problem_text
    #         i += 1
    #-----------------------------------------
    # num_problems = int(len(problem_statements))
    #------------------ Extracting the problem statements --------------------------------
    # notebook_file = input("Enter the path to the ipynb file: ")
    # programs, results, images = fix_and_execute_notebook(notebook_file)
    #----------------------------------------------------------------------
    
    
    details = {
        "{name}": name,
        "{roll_no}": roll_no,
        "{Exp_no}": exp_no,
        "{aim}": aim,
    }
    for i in range(1, num_problems + 1):
        details[f"{{{i}_problem_statement}}"] = problem_statements[i]
        details[f"{{program_{i}}}"] = programs[i-1]
        details[f"{{result_{i}}}"] = results[i] if i in results else "N/A"
    
    # Set the conclusion at the end of the document.
    details["{{conclusion}}"] = get_conclusion(aim)
    
    fill_template(template_file, output_file, details, num_problems, programs, results)
