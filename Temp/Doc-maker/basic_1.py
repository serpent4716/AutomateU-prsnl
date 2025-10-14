from docx import Document
import requests
from docx.shared import Inches, Pt 
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from notebook_extraction import *
from extraction import *
from datetime import date
today = date.today()
cached_results = {}

def get_cached_result(query: str) -> dict:
    if query not in cached_results:
        cached_results[query] = summarize_experiment(query)
    return cached_results[query]

def get_objective(query: str) -> str:
    return get_cached_result(query)["objective"]

def get_theory(query: str) -> str:
    return get_cached_result(query)["theory"]

def get_conclusion(query: str) -> str:
    return get_cached_result(query)["conclusion"]

def get_problem_statements(query: str) -> list[str]:
    return get_cached_result(query)["problem_statements",[]]

def get_references(query: str) -> list[dict]:
    return get_cached_result(query).get("references", [])
def replace_text(doc, replacements):
    """Replaces placeholders in a Word document with actual values."""
    
    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                
                para.text = para.text.replace(key, value)
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        
                        cell.text = cell.text.replace(key, value)

def set_cell_color(cell, color):
    cell_xml_element = cell._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), color)
    table_cell_properties.append(shade_obj)

def set_cell_padding(cell, top=100, start=100, bottom=100, end=100):
    """
    Sets cell margins (padding) in twips.
    Default values (100 twips ~ 0.07 inch) can be adjusted.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)

def insert_images(doc, image_replacements):
    """Inserts images into the Word document where required."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, img_path in image_replacements.items():
                    if key in cell.text:
                        cell.text = ""
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run()
                        run.add_picture(img_path, width=Inches(2.5))

def set_table_border(table):
    """Adds a border to the table by modifying its XML."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_common_rows(doc, table, i, details):
    """Adds rows common to all subjects."""
    # PROBLEM STATEMENT
    prob_row = table.add_row().cells
    prob_row[0].text = "PROBLEM STATEMENT:"
    prob_row[0].paragraphs[0].runs[0].bold = True
    prob_row[1].text = details[f"{{{i}_problem_statement}}"]
    for cell in prob_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)

    # THEORY
    theo_row = table.add_row().cells
    theo_row[0].text = "THEORY:"
    theo_row[0].paragraphs[0].runs[0].bold = True
    theo_row[1].text = details.get(f"{{theory_{i}}}", "")
    for cell in theo_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
        
    # PROGRAM
    prog_row = table.add_row().cells
    prog_row[0].text = "PROGRAM:"
    prog_row[0].paragraphs[0].runs[0].bold = True
    prog_row[1].text = details[f"{{program_{i}}}"]
    for cell in prog_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)

    # RESULT (merged)
    result_row = table.add_row().cells
    merged_result = result_row[0].merge(result_row[1])
    merged_result.text = "RESULT: {{result_%d}}" % i
    if merged_result.paragraphs and merged_result.paragraphs[0].runs:
        merged_result.paragraphs[0].runs[0].bold = True
    set_cell_padding(merged_result, top=150, start=210, bottom=150, end=150)


def add_daa_rows(doc, table, i, details):
    """Adds DAA-specific rows (theory, algo, analysis)."""

    # Pseudo-code / Algo
    algo_row = table.add_row().cells
    algo_row[0].text = "Pseudo-Code/Algo:"
    algo_row[0].paragraphs[0].runs[0].bold = True
    algo_row[1].text = details.get(f"{{algo_{i}}}", "")
    for cell in algo_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)

    # ANALYSIS
    an_row = table.add_row().cells
    an_row[0].text = "ANALYSIS:"
    an_row[0].paragraphs[0].runs[0].bold = True
    an_row[1].text = details.get(f"{{analysis_{i}}}", "")
    for cell in an_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)


def generate_subject_table(doc, subject, num_problems, details):
    for i in range(1, num_problems + 1):
        table = doc.add_table(rows=0, cols=2)

        # HEADER
        header_row = table.add_row().cells
        merged_header = header_row[0].merge(header_row[1])
        merged_header.text = f"\t\t\t\t\t\tPROGRAM {i}"
        set_cell_padding(merged_header, top=150, start=210, bottom=150, end=150)
        set_cell_color(merged_header, 'a6a6a6')

        # Common rows
        add_common_rows(doc, table, i, details)

        # Subject-specific rows
        if subject.lower() == "daa":
            add_daa_rows(doc, table, i, details)

        set_table_border(table)

def fill_template(template_path, output_path, details, num_problems, programs, results, subject):
    doc = Document(template_path)
    
    header_table = doc.add_table(rows=0, cols=2)
    name_row = header_table.add_row().cells
    name_row[0].text = "NAME"
    name_row[0].paragraphs[0].runs[0].bold = True
    name_row[1].text = "{name}"
    for cell in name_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
    
    uid_row = header_table.add_row().cells
    uid_row[0].text = "UID "
    uid_row[0].paragraphs[0].runs[0].bold = True
    uid_row[1].text = "{roll_no}"
    for cell in uid_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
    
    class_row = header_table.add_row().cells
    class_row[0].text = "Class and Batch"
    class_row[0].paragraphs[0].runs[0].bold = True
    class_row[1].text = "{class_batch}"
    for cell in class_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
        
    exp_row = header_table.add_row().cells
    exp_row[0].text = "EXP NO"
    exp_row[0].paragraphs[0].runs[0].bold = True
    exp_row[1].text = "{Exp_no}"
    for cell in exp_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
    
    date_row = header_table.add_row().cells
    date_row[0].text = "DATE"
    date_row[0].paragraphs[0].runs[0].bold = True
    date_row[1].text = str(today.strftime("%d/%m/%Y"))
    for cell in date_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
        
    aim_row = header_table.add_row().cells
    aim_row[0].text = "AIM"
    aim_row[0].paragraphs[0].runs[0].bold = True
    aim_row[1].text = "{aim}"
    for cell in aim_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
    
    obj_row = header_table.add_row().cells
    obj_row[0].text = "Objective"
    obj_row[0].paragraphs[0].runs[0].bold = True
    obj_row[1].text = "{objective}"
    for cell in obj_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
    set_table_border(header_table)
        
    generate_subject_table(doc, subject, num_problems, details)
    
    conclusion_table = doc.add_table(rows=0, cols=2)
    concl_row = conclusion_table.add_row().cells
    merged_conclusion = concl_row[0].merge(concl_row[1])
    merged_conclusion.text = "CONCLUSION: " + details["{{conclusion}}"]
    set_cell_padding(merged_conclusion, top=150, start=210, bottom=150, end=150)
    set_table_border(conclusion_table)
    
    ref_row = conclusion_table.add_row().cells
    ref_row[0].text = "References"
    ref_row[0].paragraphs[0].runs[0].bold = True
    ref_row[1].text = "{references}"
    for cell in ref_row:
        set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
    set_table_border(conclusion_table)
    
    replace_text(doc, details)
    image_placeholders = {f"{{result_{k}}}": v for k, v in results.items()}
    insert_images(doc, image_placeholders)
    
    doc.save(output_path)
    print(f"Document saved as {output_path}")

if __name__ == "__main__":
    template_file = "template.docx"  # Path to your Word template
    output_file = "output.docx"      # Output file
    
    name = input("Enter your name: ")
    roll_no = input("Enter your roll number: ")
    exp_no = input("Enter the experiment number: ")
    aim = input("Enter the aim of the experiment: ")
    subject = input("Enter the subject name: ")
    class_batch = input("Enter your class and batch: ")
    
    num_problems = int(input("Enter the number of problem statements: "))
    
    derive_from_aim = input("Do you want to derive problem statements from the aim? (yes/no): ").strip().lower()
    problem_statements = {}
    for i in range(1, num_problems + 1):
        if derive_from_aim == 'yes':
            problem_statements[i] = f"Derived from: {aim}"
        else:
            problem_statements[i] = input(f"Enter problem statement {i}: ")
    
    results = {}
    programs={}
    # for pds
    #notebook_file = input("Enter the path to the ipynb file: ")
    #programs = extract_code_from_ipynb_nbformat(notebook_file)
    for i in range (1, num_problems+1):
        programs[i] = input(f"Enter program {i}: ")
    
    details = {
        "{name}": name,
        "{roll_no}": roll_no,
        "{Exp_no}": exp_no,
        "{aim}": aim,
    }
    for i in range(1, num_problems + 1):
        details[f"{{{i}_problem_statement}}"] = problem_statements[i]
        details[f"{{program_{i}}}"] = programs[i]
        details[f"{{result_{i}}}"] = results[i] if i in results else "N/A"
    
    # Set the conclusion at the end of the document.
    details["{{conclusion}}"] = get_conclusion(aim)
    
    fill_template(template_file, output_file, details, num_problems, programs, results)
