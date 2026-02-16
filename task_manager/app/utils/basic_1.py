from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date
import json
import os
import google.generativeai as genai
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from pathlib import Path
import boto3
from io import BytesIO
from uuid import uuid4

s3 = boto3.client(
    "s3",
    region_name=os.getenv("AWS_S3_REGION"),
    aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
    aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
)

BUCKET = os.getenv("AWS_S3_BUCKET")

def upload_docx_to_s3(file_stream, filename):
    key = f"generated_docs/{uuid4()}_{filename}"

    s3.upload_fileobj(
        file_stream,
        BUCKET,
        key,
        ExtraArgs={
            "ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }
    )

    return key

BASE_DIR = Path(__file__).resolve().parent
image_path = BASE_DIR / "image.png"

from dotenv import load_dotenv
load_dotenv()


def summarize_experiment(query: str, num_statements_preference: str) -> dict:
    """
    Calls Gemini to produce structured JSON containing:
    objective, problem_statements, theory, conclusion, references
    """
    print(f"\n[INFO] Calling Gemini API for query: '{query}'...")

    problem_statement_instruction = "Generate a JSON list of relevant problem_statements (each starting with an action verb like 'To measure' or 'To analyze'). Try to minimize the number of problem statements while covering the aim comprehensively."
    if num_statements_preference == 'single':
        problem_statement_instruction = "Generate a JSON list containing exactly one problem_statement (starting with an action verb like 'To measure' or 'To analyze')."

    prompt = f"""
    You are an expert assistant that generates content for university-level technical experiment documents.

    Format your response STRICTLY as a single JSON object with the following keys:
    - "objective": A concise, single-paragraph string explaining the experiment's goal.
    - "problem_statements": {problem_statement_instruction}
    - "theory": A comprehensive and well-structured theory section of approximately 3-4 paragraphs. Explain the core concepts, relevant algorithms, and any necessary mathematical formulas. Use clear, academic language and try to beautify the text with proper formatting (like bullet points or numbered lists) where appropriate.
    - "conclusion": A 2-3 sentence string with a natural, conclusive tone.
    - "references": A JSON list of objects, where each object has a "title" string and a "uri" string.

    Input query: "{query}"
    """

    try:
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = model.generate_content(prompt)
        # Clean the response text to remove potential markdown formatting
        cleaned_text = response.text.strip().replace('```json', '').replace('```', '')
        data = json.loads(cleaned_text)
        print("[INFO] Successfully received and parsed data from Gemini.")
        return data
    except json.JSONDecodeError:
        raise ValueError("Gemini did not return valid JSON. Response was:\n" + response.text)
    except Exception as e:
        raise RuntimeError(f"An error occurred with the Gemini API call: {e}")


today = date.today()
cached_results = {}

def get_cached_result(query: str, num_statements_preference: str) -> dict:
    # Use a tuple as a key to cache based on both query and preference
    cache_key = (query, num_statements_preference)
    if cache_key not in cached_results:
        try:
            cached_results[cache_key] = summarize_experiment(query, num_statements_preference)
        except (ValueError, RuntimeError) as e:
            print(f"\n[ERROR] Could not generate content for '{query}'.")
            print(f"Details: {e}")
            # Return a default structure to prevent the script from crashing
            cached_results[cache_key] = {
                "objective": "Error: Could not generate content automatically.",
                "theory": "Error: Could not generate content automatically.",
                "conclusion": "Error: Could not generate content automatically.",
                "problem_statements": [],
                "references": []
            }
    return cached_results[cache_key]

def get_objective(query: str, pref: str) -> str:
    return get_cached_result(query, pref).get("objective", "")

def get_theory(query: str, pref: str) -> str:
    return get_cached_result(query, pref).get("theory", "")

def get_conclusion(query: str, pref: str) -> str:
    return get_cached_result(query, pref).get("conclusion", "")

def get_problem_statements(query: str, pref: str) -> list[str]:
    return get_cached_result(query, pref).get("problem_statements",[])

def get_references(query: str, pref: str) -> list[dict]:
    return get_cached_result(query, pref).get("references", [])


def set_cell_padding(cell, top=100, start=100, bottom=100, end=100):
    """Sets cell margins (padding) in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_color(cell, color):
    cell_xml_element = cell._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), color)
    table_cell_properties.append(shade_obj)

# From here you can edit the code as per your requirement
def generate_experiment_tables(doc, num_problems, details, selected_fields):
    """Generates the main table for each problem statement."""
    for i in range(1, num_problems + 1):
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        # table.autofit = False
        # --- FORMATTING CHANGE: Adjusted column widths for better presentation ---
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(6.0)

        # HEADER
        header_row = table.add_row().cells
        merged_header = header_row[0].merge(header_row[1])
        merged_header.text = f"PROGRAM {i}"
        merged_header.paragraphs[0].alignment = 1 # Center alignment
        set_cell_padding(merged_header, top=150, start=210, bottom=150, end=150)
        set_cell_color(merged_header, 'a6a6a6') # Lighter grey color

        # PROBLEM STATEMENT
        if 'Problem Statement' in selected_fields:
            prob_row = table.add_row().cells
            prob_row[0].text = "PROBLEM STATEMENT:"
            prob_row[0].paragraphs[0].runs[0].bold = True
            prob_row[1].text = details.get(f"{{problem_statement_{i}}}", "Not provided")
            for cell in prob_row: set_cell_padding(cell, top=150, start=210, bottom=150, end=150)

        
        
        # PSEUDO-CODE / ALGO
        if 'Pseudo-Code/Algo' in selected_fields:
            algo_row = table.add_row().cells
            algo_row[0].text = "Pseudo-Code/Algo:"
            algo_row[0].paragraphs[0].runs[0].bold = True
            algo_row[1].text = details.get(f"{{algo_{i}}}", "")
            for cell in algo_row: set_cell_padding(cell, top=150, start=210, bottom=150, end=150)

        # ANALYSIS
        if 'Analysis' in selected_fields:
            an_row = table.add_row().cells
            an_row[0].text = "ANALYSIS:"
            an_row[0].paragraphs[0].runs[0].bold = True
            an_row[1].text = details.get(f"{{analysis_{i}}}", "")
            for cell in an_row: set_cell_padding(cell, top=150, start=210, bottom=150, end=150)

        # PROGRAM
        if 'Program' in selected_fields:
            prog_row = table.add_row().cells
            prog_row[0].text = "PROGRAM:"
            prog_row[0].paragraphs[0].runs[0].bold = True
            prog_row[1].text = details.get(f"{{program_{i}}}", "Not provided")
            for cell in prog_row: set_cell_padding(cell, top=150, start=210, bottom=150, end=150)

        # RESULT (empty)
        if 'Result' in selected_fields:
            result_row = table.add_row().cells
            merged_result = result_row[0].merge(result_row[1])
            merged_result.text = "RESULT:"
            merged_result.paragraphs[0].runs[0].bold = True
            set_cell_padding(merged_result, top=150, start=210, bottom=150, end=150)
        
        # doc.add_paragraph() # Add space between tables
def add_college_header_stacked(doc, logo_path):
    """Adds a centered logo at the top with centered text lines below it."""
    
    # 1. Access the header
    section = doc.sections[0]
    header = section.header
    
    # Clear any existing content in the header
    for p in header.paragraphs:
        p.clear()

    # 2. ADD THE LOGO (Top Center)
    # Create a paragraph for the logo and center it
    logo_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    logo_run = logo_para.add_run()
    # Using 'r' prefix for the path string is safest for Windows
    logo_run.add_picture(str(logo_path), width=Inches(0.9)) 
    
    # 3. ADD THE TEXT (Below Logo, Center)
    header_lines = [
        "BHARATIYA VIDYA BHAVAN’S",
        "SARDAR PATEL INSTITUTE OF TECHNOLOGY",
        "Bhavan’s Campus, Munshi Nagar, Andheri (West),",
        "Mumbai – 400058-India",
        "DEPARTMENT OF COMPUTER ENGINEERING"
    ]
    
    for line in header_lines:
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(11)
        
        # Keep the lines tight together
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Add a little breathing room after the last line
    header.paragraphs[-1].paragraph_format.space_after = Pt(12)

def create_document(output_path, details, num_problems, selected_fields):
    """Creates and populates the Word document from scratch."""
    doc = Document()
    
    add_college_header_stacked(doc, logo_path=image_path)
    # --- HEADER TABLE ---
    if any(f in selected_fields for f in ['Name', 'UID', 'Class and Batch', 'Experiment No', 'Date', 'Aim', 'Objective']):
        header_table = doc.add_table(rows=0, cols=2)
        header_table.style = 'Table Grid'
        # --- FORMATTING CHANGE: Set column widths to match desired layout ---
        header_table.columns[0].width = Inches(1.5)
        header_table.columns[1].width = Inches(6.0)

        if 'Name' in selected_fields:
            name_row = header_table.add_row().cells
            name_row[0].text = "NAME"
            name_row[0].paragraphs[0].runs[0].bold = True
            name_row[1].text = details.get('{name}', '')
        
        if 'UID' in selected_fields:
            uid_row = header_table.add_row().cells
            uid_row[0].text = "UID"
            uid_row[0].paragraphs[0].runs[0].bold = True
            uid_row[1].text = details.get('{uid}', '')
        
        if 'Class and Batch' in selected_fields:
            class_row = header_table.add_row().cells
            class_row[0].text = "Class and Batch"
            class_row[0].paragraphs[0].runs[0].bold = True
            class_row[1].text = details.get('{class_batch}', '')
            
        if 'Experiment No' in selected_fields:
            exp_row = header_table.add_row().cells
            exp_row[0].text = "EXP NO"
            exp_row[0].paragraphs[0].runs[0].bold = True
            exp_row[1].text = details.get('{exp_no}', '')
        
        if 'Date' in selected_fields:
            date_row = header_table.add_row().cells
            date_row[0].text = "DATE"
            date_row[0].paragraphs[0].runs[0].bold = True
            date_row[1].text = details.get('{date}', '')
        
        if 'Aim' in selected_fields:
            aim_row = header_table.add_row().cells
            aim_row[0].text = "AIM"
            aim_row[0].paragraphs[0].runs[0].bold = True
            aim_row[1].text = details.get('{aim}', '')
        
        if 'Objective' in selected_fields:
            obj_row = header_table.add_row().cells
            obj_row[0].text = "Objective"
            obj_row[0].paragraphs[0].runs[0].bold = True
            obj_row[1].text = details.get('{objective}', '')

        # THEORY
        if 'Theory' in selected_fields:
            theo_row = header_table.add_row().cells
            theo_row[0].text = "THEORY:"
            theo_row[0].paragraphs[0].runs[0].bold = True
            theo_row[1].text = details.get("{theory}", "")
            

        for row in header_table.rows:
            for cell in row.cells:
                set_cell_padding(cell, top=150, start=210, bottom=150, end=150)
        # doc.add_paragraph()

    # --- EXPERIMENT DETAILS TABLES ---
    if num_problems > 0:
        generate_experiment_tables(doc, num_problems, details, selected_fields)
    
    # --- CONCLUSION AND REFERENCES TABLE ---
    if any(f in selected_fields for f in ['Conclusion', 'References']):
        conclusion_table = doc.add_table(rows=0, cols=1)
        conclusion_table.style = 'Table Grid'
        conclusion_table.columns[0].width = Inches(7.5)


        if 'Conclusion' in selected_fields:
            concl_row = conclusion_table.add_row().cells
            paragraph = concl_row[0].paragraphs[0]
            run = paragraph.add_run("CONCLUSION:")
            run.bold = True
            paragraph.add_run("\n\t" + details.get("{{conclusion}}", ""))
            
            
            
        if 'References' in selected_fields:
            ref_row = conclusion_table.add_row().cells
            paragraph = ref_row[0].paragraphs[0]
            run = paragraph.add_run("REFERENCES:")
            run.bold = True
            paragraph.add_run("\n")
            references = details.get("{references}", [])
            if references:
                for ref in references:
                    title = ref.get('title', 'Untitled')
                    uri = ref.get('uri', ref.get('url', 'N/A'))
                    # Each reference on a new line with an indent and bullet
                    paragraph.add_run(f"\t• {title}: {uri}\n")
            else:
                paragraph.add_run("\tNo references generated.")
        
        for row in conclusion_table.rows:
            for cell in row.cells:
                set_cell_padding(cell, top=150, start=210, bottom=150, end=150)

    # doc.save(output_path)
    # print(f"\nDocument successfully saved as {output_path}")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

def get_user_choices():
    """Gets user selection for which fields to include in the document."""
    fields = ['Name', 'UID', 'Class and Batch', 'Experiment No', 'Date', 'Aim', 'Objective', 'Problem Statement', 'Theory', 'Pseudo-Code/Algo', 'Analysis', 'Program', 'Result', 'Conclusion', 'References']
    print("Please choose the fields to include in your document.")
    print("Enter the numbers separated by commas (e.g., 1,2,8,12).")
    for i, field in enumerate(fields, 1):
        print(f"  {i}. {field}")
    
    while True:
        try:
            choices = input("> ")
            selected_indices = [int(c.strip()) - 1 for c in choices.split(',')]
            selected_fields = {fields[i] for i in selected_indices if 0 <= i < len(fields)}
            if not selected_fields:
                print("Please select at least one field.")
                continue
            print("\nYou have selected:", ", ".join(sorted(list(selected_fields))))
            return selected_fields
        except (ValueError, IndexError):
            print("Invalid input. Please enter numbers from the list, separated by commas.")


def run_full_generation_process(payload_dict, task_id):
    """
    This is the main "business logic" function.
    It's called by Celery but knows nothing about the DB.
    It does all the work and returns the final file path or raises
    an exception.
    """
    
    print(f"[Task {task_id}] Generation process started.")
    
    # We wrap the *entire* logic in a try...finally
    # to ensure we don't leave any partial work.
    try:
        # 2. Parse payload
        basic_details = payload_dict.get('basicDetails', {})
        selected_fields = payload_dict.get('selectedSections', [])
        ps_pref = payload_dict.get('problemStatementCount', 'single')
        aim = basic_details.get('Aim')

        if not aim:
            raise ValueError("Aim is missing from basicDetails")

        # 3. Build 'details' dict
        details = {}
        # This is the same logic that *was* in your tasks.py
        if 'Name' in selected_fields:
            details['{name}'] = basic_details.get('Name', '')
        if 'UID' in selected_fields:
            details['{uid}'] = basic_details.get('UID', '')
        if 'Class and Batch' in selected_fields:
            details['{class_batch}'] = basic_details.get('Class_and_Batch', '')
        if 'Experiment No' in selected_fields:
            details['{exp_no}'] = basic_details.get('Experiment_No', '')
        if 'Date' in selected_fields:
            details['{date}'] = basic_details.get('Date', today.strftime("%d/%m/%Y"))
        if 'Aim' in selected_fields:
            details['{aim}'] = aim
        if 'Objective' in selected_fields:
            details['{objective}'] = get_objective(aim, ps_pref)
        if 'Theory' in selected_fields:
            details['{theory}'] = get_theory(aim, ps_pref)
        if 'Conclusion' in selected_fields:
            details['{{conclusion}}'] = get_conclusion(aim, ps_pref)
        if 'References' in selected_fields:
            details['{references}'] = get_references(aim, ps_pref)
        
        problem_statements = []
        if 'Problem Statement' in selected_fields:
            problem_statements = get_problem_statements(aim, ps_pref)
        
        num_problems = len(problem_statements)

        # Handle problem-specific details
        for i in range(1, num_problems + 1):
            problem = problem_statements[i-1]
            
            details[f"{{problem_statement_{i}}}"] = problem
            
            if 'Program' in selected_fields:
                details[f"{{program_{i}}}"] = ""
            if 'Pseudo-Code/Algo' in selected_fields:
                details[f"{{algo_{i}}}"] = ""
            if 'Analysis' in selected_fields:
                details[f"{{analysis_{i}}}"] = ""
            if 'Result' in selected_fields:
                details[f"{{result_{i}}}"] = ""
        
        # 5. Call your create_document function
        # create_document(
        #     output_path=output_file_path,
        #     details=details,
        #     num_problems=num_problems,
        #     selected_fields=selected_fields
        # )
        
        # # 6. Return the final path on success
        # print(f"[Task {task_id}] Generation process complete.")
        # return output_file_path
        file_stream = create_document(
                details=details,
                num_problems=num_problems,
                selected_fields=selected_fields
            )

        url = upload_docx_to_s3(file_stream, f"{task_id}.docx")

        print(f"[Task {task_id}] Uploaded to S3: {url}")
        return url
        
    except Exception as e:
        # 7. If anything fails, print and re-raise the exception
        # so the Celery task can catch it and update the DB.
        print(f"[Task {task_id}] ERROR in run_full_generation_process: {e}")
        # Optional: Clean up failed file
        # if os.path.exists(output_file_path):
        #     os.remove(output_file_path)
        raise e
    

if __name__ == "__main__":
    output_file = "experiment.docx"
    
    selected_fields = get_user_choices()

    print("\nPlease provide the following details for your experiment:")
    details = {}
    if 'Name' in selected_fields: details['{name}'] = input("Enter your name: ")
    if 'UID' in selected_fields: details['{uid}'] = input("Enter your UID: ")
    if 'Class and Batch' in selected_fields: details['{class_batch}'] = input("Enter your class and batch: ")
    if 'Experiment No' in selected_fields: details['{exp_no}'] = input("Enter the experiment number: ")
    
    # Either get a pdf as input or just copy pasted text
    # choice = input("Do you want to (1) paste the AIM or (2) provide a PDF file path? Enter 1 or 2: ").strip()
    # if choice == '2':
    #     pdf_path = input("Enter the full path to the PDF file: ").strip()
    #     try:
    #         with open(pdf_path, 'rb') as pdf_file:
    #             pdf_bytes = pdf_file.read()
    #             aim = extraction.process_document(pdf_bytes)
    #             print(f"\nExtracted AIM from PDF:\n{aim}\n")
    #     except Exception as e:
    #         print(f"Error reading PDF file: {e}")
    #         aim = input("Please enter the AIM manually: ")
    # elif choice == '1':
    #     aim = input("Enter the AIM of the experiment: ")
    # else:
    #     print("Invalid choice. Continuing with manual AIM entry.")
    aim = input("Enter the main AIM of the experiment: ")
    details['{aim}'] = aim

    # --- NEW FEATURE: Ask user for problem statement preference ---
    while True:
        ps_pref = input("Do you want a 'single' or 'multiple' problem statements? ").lower().strip()
        if ps_pref in ['single', 'multiple']:
            break
        print("Invalid choice. Please enter 'single' or 'multiple'.")

    # Call the summarize function to get automated content
    details['{objective}'] = get_objective(aim, ps_pref)
    details['{theory}'] = get_theory(aim, ps_pref)
    details['{{conclusion}}'] = get_conclusion(aim, ps_pref)
    details['{references}'] = get_references(aim, ps_pref)
    details['{date}'] = today.strftime("%d/%m/%Y")
    
    problem_statements = get_problem_statements(aim, ps_pref)
    num_problems = len(problem_statements)

    if num_problems == 0:
        print("\n[WARNING] Could not automatically derive problem statements from the aim.")
        try:
            num_problems = int(input("Please enter the number of problem statements manually: "))
        except ValueError:
            print("Invalid number. Exiting.")
            num_problems = 0
    else:
        print(f"Found {num_problems} problem statement(s) based on the aim.")

    for i in range(1, num_problems + 1):
        print(f"\n--- Details for Program {i} ---")
        
        if problem_statements:
            details[f"{{problem_statement_{i}}}"] = problem_statements[i-1]
            print(f"Problem Statement: {problem_statements[i-1]}")
        elif 'Problem Statement' in selected_fields:
             details[f"{{problem_statement_{i}}}"] = input(f"Enter Problem Statement for Program {i}:\n")

        if 'Program' in selected_fields: details[f"{{program_{i}}}"] = input(f"Enter or paste the code for Program {i}:\n")
        if 'Result' in selected_fields: details[f"{{result_{i}}}"] = ""  # Placeholder for Result
        if 'Pseudo-Code/Algo' in selected_fields: details[f"{{algo_{i}}}"] = input(f"Enter Pseudo-Code/Algo for Program {i}:\n")
        if 'Analysis' in selected_fields: details[f"{{analysis_{i}}}"] = input(f"Enter Analysis for Program {i}:\n")

    create_document(output_file, details, num_problems, selected_fields)

